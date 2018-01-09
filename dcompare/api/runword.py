"""
Author: chiyuan
Date created: 12/10/17
python version: 3.6

performs the comparison logic based on the output from compare.py and outputs summary.docx.
"""

from .compare2 import DocCompare
from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml 
import pdb
import difflib
import re
import logging


stop_words = ['in', 'the', 'of', 'or', 'to', 'with', 'and', 'on', 'without']

def process_paragraph(raw_text):
    text = re.sub(r'(\xa0+)', '\t', raw_text)
    text = re.sub(r'(\s\s+)', '\t', text)
    if len(text.split("\t")) > 1:
        leave_subsec = text.split("\t")[1].strip()
    else:
        leave_subsec = text.strip()
        
    return leave_subsec

def remove_stopwords(s):
    no_stop = [f for f in s.split(" ") if f not in stop_words]
    return " ".join(no_stop)
    
def find_doc_clause(clause, cur_clauses, clause_idx, document, sect_name, sections, sect_idx, tol=.7):
    """
    Given a clause (from 1 document), return the same clause in from another document
    :return: (clause text, clause name, section name)
    """
    #cluase in current section, skip
    if clause in document[sect_name]:
        return document[sect_name][clause], clause, sect_name
    
    #doc in current
    for s in list(sections)[sect_idx + 1:]:
        if clause in document[s]:
            return document[s][clause], clause, s

    #edit distance betewen ground truth clause and prospecticve clauses. match max score clause if
    #above certain clause
    ratios = []
    for s in list(sections)[sect_idx : ]:
        for clause2 in document[s].keys():
            if clause2 in cur_clauses[clause_idx + 1: ]:
                tup = (s, difflib.SequenceMatcher(None, clause, clause2).ratio(), clause2)
                ratios.append(tup)
    
    if len(ratios) == 0:
        return None
        
    max_sect, max_ratio, max_clause = max(ratios, key = lambda i: i[1])
    if max_ratio >= tol: #TODO: maybe to add in additional check here
        return document[max_sect][max_clause], max_clause, max_sect# == sect_name
    else:
        return None

def style_diffs(para, text, opcodes, is_a):
    for code, a0, a1, b0, b1 in opcodes:
        if code == 'equal':
            if is_a:
                s,e = a0, a1
            else:
                s,e = b0, b1
            para.add_run(text[s:e])
        elif code == 'insert':
            if not is_a:
                para.add_run(text[b0:b1]).font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
        elif code == 'delete':
            if is_a:
                para.add_run(text[a0:a1]).font.highlight_color = WD_COLOR_INDEX.PINK
        elif code == 'replace':
            if is_a:
                para.add_run(text[a0:a1]).font.highlight_color = WD_COLOR_INDEX.TURQUOISE
            else:
                para.add_run(text[b0:b1]).font.highlight_color = WD_COLOR_INDEX.TURQUOISE
        
        
        
def run2(files, skip_nochange = True):
    #TODO: ideally this should yield {["section name": [{"row name": ["text1", "text2", "text3"]}]]
    """
    Compares tender documents and groups them according to sections and clauses in a table format.
    """
    content_key = 'content'

    dcom = DocCompare()
    for file1 in files:
        dcom.add_doc(file1)

    contents = dcom.all_contents()

    # set content rows
    bigsection_names = contents[0].keys()

    # each section
    for bs_idx, bs_name in enumerate(bigsection_names):
        row_headers = list(set([k for d in contents for k in d[bs_name].keys()]))
        exclude = []

        section_contents = []

        # each clause in current section
        for header_idx, header in enumerate(row_headers):
            if header in exclude:
                continue

            sim = []

            # cur_row = table.add_row().cells
            # cur_row[0].text = header

            matched_clauses = [find_doc_clause(header, row_headers, header_idx, doc, bs_name, bigsection_names, bs_idx) \
                               for doc_idx, doc in enumerate(contents)]

            hashes = set([m[0]['hash'] if m is not None else 'hello world' for m in matched_clauses])

            compare_indices = []
            added_hashes = []
            seqm = None
            oc = None

            #this is to change there are 2 distinct columns in this row.
            if len(hashes) == 2 and 'hello world' not in hashes:
                for i, (dcontent, dclause, dsect) in enumerate(matched_clauses):
                    if dcontent['hash'] in hashes and dcontent['hash'] not in added_hashes:
                        compare_indices.append(i)
                        added_hashes.append(dcontent['hash'])

                seqm = difflib.SequenceMatcher(None,
                                               "\n\n".join([process_paragraph(para.text) for para in
                                                            matched_clauses[compare_indices[0]][0][content_key]]),
                                               "\n\n".join([process_paragraph(para.text) for para in
                                                            matched_clauses[compare_indices[1]][0][content_key]])
                                               )
                oc = seqm.get_opcodes()
                assert len(added_hashes) == 2

            c1 = 0
            c2 = 0
            row_text = []

            #for each document what is the matched clause in that document
            for i, m in enumerate(matched_clauses):
                org_name = ''
                org_section = ''
                t = ''
                if m is not None:
                    clause_content, clause_name, sect = m
                    # changed = False
                    if bs_name == sect:
                        exclude.append(clause_name)
                    else:
                        contents[i][sect].pop(clause_name)
                        org_section = sect
                        # t = t + "This clause is originally in {}\n".format(sect)
                        # changed = True

                    if clause_name != header:
                        org_name = clause_name
                        # t = t + "This clause is originally named {}\n".format(clause_name)
                        # changed = True

                    # if changed:
                    #     t = t + "---------------------------------------------------------------\n\n"

                if len(hashes) == 1:
                    t += "NO CHANGES"

                # elif oc is not None and (c1 == 0 or c2 == 0): #display only 2 columns if there are 2 unique columns
                #     clause_content, clause_name, sect = m
                #     # p = cur_row[i + 1].add_paragraph()
                #     # p.add_run(t)
                #     if clause_content['hash'] == added_hashes[0]:
                #         # style_diffs(p, seqm.a, oc, True)
                #         c1 += 1
                #     else:
                #         # style_diffs(p, seqm.b, oc, False)
                #         c2 += 1
                #     #continue

                elif m is not None:
                    t = t + "\n\n".join([para.text for para in clause_content[content_key]])

                # cur_row[i + 1].text = t
                # row_text.append(t)
                row_text.append({'text': t, 'org_name': org_name, 'org_section': org_section})

            section_contents.append({'section_name':header, 'section': row_text})

        yield {'bs_name':bs_name, 'sections': section_contents}


def run(files, skip_nochange=True):
    """
    Compares tender documents and groups them according to sections and clauses in a table format.
    Outputs a docx document.
    """
    #file1 = 'itt_edited.docx'
    #file2 = 'itt_edited_modified.docx'

    #file1 = "C:/Users/mas_cygoh/Desktop/stressfulshit/stressfulshit/CCOM3/1.docx"
    file1 = "C:/Users/mas_cygoh/Desktop/docs/docs/Consolidated Tender Docs/2.docx"
    file2 = "C:/Users/mas_cygoh/Desktop/docs/docs/CCOM3/1.docx"
    file3 = "C:/Users/mas_cygoh/Desktop/docs/docs/3.docx"
    
    sect_colors = ['7BA65D', 'B6D6F2', 'D9A9BF', 'D9D9D9']
    
    dcom = DocCompare()
    for file1 in files:
        dcom.add_doc(file1)
    #dcom.add_doc(file2)
    #dcom.add_doc(file3)
    #dcom.add_doc(file2)
    #dcom.add_doc(file2)
    contents = dcom.all_contents()

    summ = Document()
    section = summ.sections[-1]
    w,h = section.page_height, section.page_width
    section.page_width = w
    section.page_height = h
    section.left_margin = Inches(0.15)
    section.right_margin = Inches(0.15)
    
    """
    nrows = []
    for content in contents: #each doc
        for k in content.values(): #each big section
            org = []
            for k2 in k.keys(): #each section
                org.append(k2)
                    
            nrows += list(set(org))
    nrows = nrows
    """

    table = summ.add_table(rows=0, cols=len(contents) + 1)
    table.style = 'TableGrid'
    header_row = table.add_row().cells#table.rows[0].cells
    #set header row
    for c in range(len(contents)):
        para = header_row[c + 1].add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run(dcom.documents[c]).bold = True
    
    #set content rows
    bigsection_names = contents[0].keys()
    
    count = 0
    
    #each section
    for bs_idx, bs_name in enumerate(bigsection_names):
        row = table.add_row()
        fcell = row.cells[0]
        lcell = row.cells[len(contents)]
        m = fcell.merge(lcell)
        para = m.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run(bs_name).bold = True
        
        #add background colors for section header cells
        shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), sect_colors[bs_idx%len(sect_colors)]))
        m._tc.get_or_add_tcPr().append(shading_elm)
                
        row_headers = list(set([k for d in contents for k in d[bs_name].keys()]))
        exclude = []
        
        #each clause in current section
        for header_idx, header in enumerate(row_headers):
            if header in exclude:
                continue
                
            sim = [] 
            
            """
            for j, doc_content in enumerate(contents):
                d = doc_content[bs_name].get(header, {})
                sim.append( d.get("hash", "helloworld") )
            """
            
            cur_row = table.add_row().cells
            cur_row[0].text = header
            
            """
            if len(set(sim)) == 1 and skip_nochange:
                for j in range(len(contents)):
                    cur_row[j + 1].text = "NO CHANGE"
                continue
            """
            
            matched_clauses = [find_doc_clause(header, row_headers, header_idx, doc, bs_name, bigsection_names, bs_idx) \
            for doc_idx, doc in enumerate(contents)]
            
            hashes = set([m[0]['hash'] if m is not None else 'hello world' for m in matched_clauses])
            
            compare_indices = []
            added_hashes = []
            seqm = None
            oc = None
            
            if len(hashes) == 2 and 'hello world' not in hashes:    
                for i, (dcontent, dclause, dsect) in enumerate(matched_clauses):
                    if dcontent['hash'] in hashes and dcontent['hash'] not in added_hashes:
                        compare_indices.append(i)
                        added_hashes.append(dcontent['hash'])
                   
                seqm = difflib.SequenceMatcher(None, 
                    "\n\n".join([process_paragraph(para.text) for para in matched_clauses[compare_indices[0]][0]['section']]),
                    "\n\n".join([process_paragraph(para.text) for para in matched_clauses[compare_indices[1]][0]['section']])  
                )
                oc = seqm.get_opcodes()                
                assert len(added_hashes) == 2
                
                #if header == 'acceptance of tender proposal':
                #    pdb.set_trace()
            
            c1 = 0
            c2 = 0
            for i, m in enumerate(matched_clauses):
                t = ''
                if m is not None:
                    clause_content, clause_name, sect = m
                    changed = False
                    if bs_name == sect:
                        exclude.append(clause_name)
                    else:
                        contents[i][sect].pop(clause_name)
                        t = t + "This clause is originally in {}\n".format(sect)
                        changed = True
                    
                    if clause_name != header:
                        t = t + "This clause is originally named {}\n".format(clause_name)
                        changed = True
                    
                    if changed:
                        t = t + "---------------------------------------------------------------\n\n"
                
                if len(hashes) == 1:
                    t += "NO CHANGES"
                elif oc is not None and (c1 == 0 or c2 == 0):
                    clause_content, clause_name, sect = m
                    p = cur_row[i + 1].add_paragraph()
                    p.add_run(t)
                    if clause_content['hash'] == added_hashes[0]:
                        style_diffs(p, seqm.a, oc, True)
                        c1 += 1
                    else:
                        style_diffs(p, seqm.b, oc, False)
                        c2 += 1
                    continue
                    
                elif m is not None:
                    t = t + "\n\n".join([para.text for para in clause_content['section']])
                    
                cur_row[i + 1].text = t
            
            """
            for doc_idx, doc_content in enumerate(contents):
                d = doc_content[bs_name].get(header, {})
                
                #can't find the clause under the current section
                #search in other sections
                if not any(d): 
                    candidates = []
                    ratios = []
                    for bs_name2,v in doc_content.items():
                        if header in v and bs_name2 != bs_name:
                            candidates.append((bs_name2, v[header]) )
                            
                        for header2 in v.keys():
                            r = difflib.SequenceMatcher(None, remove_stopwords(header2), remove_stopwords(header)).ratio()
                            ratios.append( (r, header2, bs_name2) )
                    
                    if len(candidates) == 1:
                        sectname, d = candidates[0]
                        paras = d.get("section", [])
                        cur_row[doc_idx + 1].text = "This clause is originally in %s.\n"%sectname + \
                                "======================================\n\n" + "\n\n".join([para.text for para in paras])
                        contents[doc_idx][sectname].pop(header, None)
                        #pdb.set_trace()
                        continue
                    
                    else:
                        sratio = sorted(ratios, key=lambda x: x[0])
                        max_ratio, max_header, max_section = sratio[-1]
                        if max_ratio > 0.6:
                            t = "\n\n".join([para.text for para in paras])
                            if max_section != bs_name:
                                cur_row[doc_idx + 1].text = "This clause is originally in %s named %s.\n"%(max_section, max_header)  + \
                                "======================================\n\n" + t
                            else:
                                cur_row[doc_idx + 1].text = "This clause is differently named %s\n"%(max_header)  + \
                                "======================================\n\n" + t
                            
                            contents[doc_idx][max_section].pop(max_header, None)
                            continue
                    
                    
                    
                paras = d.get("section", [])
                cur_row[doc_idx + 1].text = "\n\n".join([para.text for para in paras])
                """
    """         
    for header in nrows:
        cur_row = table.rows[r].cells
        cur_row[0].text = header
        for j, doc_content in enumerate(contents):
            d = doc_content.get(header, {})
            paras = d.get("section", "")
            cur_row[j + 1].text = "\n".join([para.text for para in paras])
        r += 1
    """
    print("total count: {}".format(count))
    summ.save("summary.docx")

if __name__ == '__main__':
    run()

